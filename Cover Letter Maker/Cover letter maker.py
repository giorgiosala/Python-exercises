#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jul 26 21:28:50 2021

@author: giorgiosala
"""

from docx import *

import random
import os

df={}

def setup():
        profile=input("Add new job Profile? Y/N")
        if profile=="y":
           job_title=input("Insert Name of Job title: ")
           expertise=input("Insert Name of Your expertise: ")
           hardskill1=input("Insert Name of Hard Skill 1:")
           hardskill2=input("Insert Name of Hard Skill 2:")
           hardskill3=input("Insert Name of Hard Skill 2:")
           
        else :
            pass
        
        return job_title
   
setup()
print(job_title)
   
    
"""
document = Document()

hm=input("Insert Name of Hiring Manager: ")
company=input("Insert Name of Company: ")
job_title=input("Insert Name of Job title: ")
expertise=input("Insert Name of Your expertise: ")

companylike=input("What you like about the company: ")
likingcompanymission=input("Whant you like about the company mission: ")
document.add_paragraph(f"Dear {hm},")
document.add_paragraph(f"I’m incredibly excited to submit my application for {job_title} at {company}")
document.add_paragraph(f"As a {expertise} professional, I have experience in Life sciences, Data Analisis and Team Management. In my previous roles, I have driven my project to success through:")
document.add_paragraph('Dedication to develop and perform my protocols, which led to my Master thesis', style='List Bullet')
document.add_paragraph('Always studying and working on how to express my data in the most clear and detailed way', style='List Bullet')
document.add_paragraph('Focused on little details', style='List Bullet')
document.add_paragraph(f"{company} is of particular interest to me given {companylike}.{likingcompanymission}")
document.add_paragraph(f"I would love to have a moment to discuss on how my particular skills set will bring value to your organization,")
document.add_paragraph(f"Thank you for the consideration")

document.save('/Users/giorgiosala/Desktop/demo.docx')"""

        
        



